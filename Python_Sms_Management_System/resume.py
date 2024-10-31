# Creating the resume for MERN Stack Developer
from docx import Document

# Resume data for MERN Stack Developer
mern_stack_data = {
    "role": "MERN Stack Developer",
    "summary": (
        "Detail-oriented MERN Stack Developer with extensive experience in designing, "
        "developing, and deploying full-stack web applications. Proficient in using "
        "MongoDB, Express.js, React.js, and Node.js to create responsive and high-performance "
        "web applications. Strong understanding of the software development lifecycle, "
        "team collaboration, and problem-solving skills."
    ),
    "experience": [
        {
            "title": "MERN Stack Developer",
            "company": "Freelance",
            "duration": "Jan 2023 – Present",
            "description": [
                "Designed and developed a full-stack e-commerce application using the MERN stack, "
                "implementing features such as user authentication, product management, and payment processing.",
                "Collaborated with clients to gather requirements and deliver customized web solutions, "
                "enhancing user experience and client satisfaction.",
                "Optimized application performance by implementing efficient database queries and caching strategies.",
                "Integrated third-party APIs for payment gateways and shipping services, ensuring seamless operations."
            ]
        },
        {
            "title": "Web Development Intern",
            "company": "Interns.pk",
            "duration": "Feb 2024 – March 2024",
            "description": [
                "Assisted in the development of a multi-user blogging platform using the MERN stack, "
                "contributing to both frontend and backend functionalities.",
                "Participated in code reviews and contributed to optimizing existing codebases, improving application speed.",
                "Conducted user testing sessions to gather feedback and iterate on features, enhancing usability."
            ]
        }
    ],
    "education": "Bachelor of Science in Chemistry, Federal University of Ebonyi State | 2020 – Present\nWeb Development, Floritech Computer College | April 2021 – October 2021",
    "skills": [
        "Frontend: React.js, HTML5, CSS3, JavaScript",
        "Backend: Node.js, Express.js, MongoDB",
        "Tools: Git/GitHub, Postman, NPM, Webpack, Docker",
        "Other Skills: RESTful API Development, JWT Authentication, Redux for State Management"
    ],
    "projects": [
        {
            "title": "E-commerce Platform",
            "technologies": "React.js, Node.js, Express.js, MongoDB",
            "description": "Developed a comprehensive e-commerce platform with user authentication, shopping cart, and payment integration."
        },
        {
            "title": "Task Management Application",
            "technologies": "MongoDB, Express.js, React.js, Node.js",
            "description": "Created a task management tool featuring real-time collaboration and customizable task organization."
        }
    ],
    "certifications": [
        "MERN Stack Web Development – Udemy, January 2024",
        "JavaScript Algorithms and Data Structures – FreeCodeCamp, June 2023"
    ],
    "languages": "English (Fluent)"
}

# Function to create the resume document for a specific role
def create_resume_doc(role_data):
    doc = Document()
    doc.add_heading(f"{role_data['role']} Resume", 0)

    # Adding personal information
    doc.add_paragraph("Chibuike Praise Obiora\n61, Adediran Street, Sabo - Oniba, Ojo, Lagos State\nZip Code: 102115 | Phone Number: 08025160310\nEmail: obiorachibuike22@gmail.com | LinkedIn: linkedin.com/in/obiora-chibuike\nGitHub: github.com/obiorachibuike | Website: codetech-pi.vercel.app")

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(role_data['summary'])

    # Experience
    doc.add_heading('Experience', level=1)
    for experience in role_data['experience']:
        doc.add_heading(experience['title'], level=2)
        doc.add_paragraph(f"{experience['company']} | {experience['duration']}")
        for bullet in experience['description']:
            doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph(role_data['education'])

    # Skills
    doc.add_heading('Skills', level=1)
    for skill in role_data['skills']:
        doc.add_paragraph(skill)

    # Projects
    doc.add_heading('Projects', level=1)
    for project in role_data['projects']:
        doc.add_heading(project['title'], level=2)
        doc.add_paragraph(f"Technologies: {project['technologies']}")
        doc.add_paragraph(project['description'])

    # Certifications
    doc.add_heading('Certifications', level=1)
    for certification in role_data['certifications']:
        doc.add_paragraph(certification)

    # Languages
    doc.add_heading('Languages', level=1)
    doc.add_paragraph(role_data['languages'])

    return doc

# Create the MERN Stack Developer document
mern_doc = create_resume_doc(mern_stack_data)
mern_doc_path = "/mnt/data/MERN_Stack_Developer_Resume.docx"
mern_doc.save(mern_doc_path)

mern_doc_path
